"""DOCX questionnaire parser.

Walks the document body in true reading order and produces a
:class:`~app.models.document.ParsedDocument`: paragraphs, tables, question
boundaries and character-level bold/italic formatting.

``python-docx`` exposes ``Document.paragraphs`` and ``Document.tables`` as two
separate sequences, which loses the interleaving between them. A questionnaire
depends on that interleaving — a grid's table has to stay attached to the
question stem above it — so the body XML is walked directly instead.
"""

from __future__ import annotations

import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import BinaryIO

import docx
from docx.oxml.ns import qn
from docx.table import Table as DocxTable, _Cell
from docx.text.paragraph import Paragraph as DocxParagraph

from app.models.document import (
    Block,
    DocumentStats,
    ListInfo,
    ParagraphBlock,
    ParsedDocument,
    TableBlock,
    TableCell,
    TableRow,
)
from app.parsing.formatting import document_defaults, extract_runs, runs_to_text
from app.parsing.normalize import collapse_whitespace, detect_literal_marker
from app.parsing.numbering import NumberingResolver, paragraph_numbering
from app.parsing.question_boundaries import BoundaryConfig, detect_boundaries


class DocxParseError(ValueError):
    """Raised when a file cannot be read as a Word document."""


@dataclass
class _ParseContext:
    document: object
    defaults: dict[str, bool]
    numbering: NumberingResolver
    stats: DocumentStats = field(default_factory=DocumentStats)
    _next_index: int = 0

    def take_index(self) -> int:
        index = self._next_index
        self._next_index += 1
        return index


def parse_docx(
    source: str | Path | BinaryIO,
    filename: str | None = None,
    boundary_config: BoundaryConfig | None = None,
) -> ParsedDocument:
    """Parse a DOCX file from a path or an open binary stream."""
    if isinstance(source, (str, Path)):
        path = Path(source)
        if not path.is_file():
            raise DocxParseError(f"No such file: {path}")
        filename = filename or path.name

    try:
        document = docx.Document(source)
    except (zipfile.BadZipFile, KeyError, ValueError) as exc:
        raise DocxParseError(
            "File could not be read as a .docx document. Legacy .doc files and "
            "renamed files are not supported."
        ) from exc

    context = _ParseContext(
        document=document,
        defaults=document_defaults(document),
        numbering=NumberingResolver(document),
    )

    blocks = _walk_container(document.element.body, document, context)
    questions, warnings = detect_boundaries(blocks, boundary_config)

    context.stats.questions = sum(1 for q in questions if not q.is_preamble)

    return ParsedDocument(
        source_filename=filename,
        blocks=blocks,
        questions=questions,
        stats=context.stats,
        warnings=warnings,
    )


# -- body traversal -------------------------------------------------------


def _walk_container(element, parent, context: _ParseContext, depth: int = 0) -> list[Block]:
    """Collect blocks from a body, cell or content-control in document order."""
    blocks: list[Block] = []
    for child in element:
        tag = child.tag
        if tag == qn("w:p"):
            blocks.append(_build_paragraph(DocxParagraph(child, parent), context))
        elif tag == qn("w:tbl"):
            blocks.append(_build_table(DocxTable(child, parent), context, depth))
        elif tag == qn("w:sdt"):
            # A content control wraps its real content one level down.
            content = child.find(qn("w:sdtContent"))
            if content is not None:
                blocks.extend(_walk_container(content, parent, context, depth))
    return blocks


def _build_paragraph(paragraph: DocxParagraph, context: _ParseContext) -> ParagraphBlock:
    runs = extract_runs(paragraph, context.defaults)
    text = collapse_whitespace(runs_to_text(runs))

    context.stats.paragraphs += 1
    context.stats.runs += len(runs)
    context.stats.bold_runs += sum(1 for r in runs if r.bold)
    context.stats.italic_runs += sum(1 for r in runs if r.italic)
    if text:
        context.stats.non_empty_paragraphs += 1

    style_name = _style_name(paragraph)
    return ParagraphBlock(
        index=context.take_index(),
        text=text,
        runs=runs,
        style=style_name,
        heading_level=_heading_level(style_name, paragraph),
        alignment=paragraph.alignment.name if paragraph.alignment is not None else None,
        list_info=_list_info(paragraph, context),
        literal_marker=detect_literal_marker(text),
    )


def _build_table(table: DocxTable, context: _ParseContext, depth: int) -> TableBlock:
    context.stats.tables += 1
    # Claim the table's index before walking its cells so that a parent block
    # always sorts ahead of its nested content.
    table_index = context.take_index()
    rows: list[TableRow] = []
    max_columns = 0

    for row_index, row in enumerate(table.rows):
        cells: list[TableCell] = []
        column = 0
        for cell in _iter_row_cells(row):
            grid_span = _grid_span(cell)
            v_merge = _vertical_merge(cell)
            cell_blocks = _walk_container(cell._tc, cell, context, depth + 1)
            cells.append(
                TableCell(
                    row=row_index,
                    col=column,
                    grid_span=grid_span,
                    v_merge=v_merge,
                    blocks=cell_blocks,
                    text=_blocks_to_text(cell_blocks),
                )
            )
            column += grid_span
        max_columns = max(max_columns, column)
        rows.append(TableRow(index=row_index, cells=cells, is_header=_is_header_row(row)))

    return TableBlock(
        index=table_index,
        n_rows=len(rows),
        n_cols=max_columns,
        rows=rows,
        style=_table_style_name(table),
        nesting_depth=depth,
    )


def _iter_row_cells(row):
    """Yield each ``w:tc`` of a row once.

    ``Row.cells`` repeats a horizontally merged cell once per grid column,
    which would duplicate its text; the underlying elements do not.
    """
    for tc in row._tr.findall(qn("w:tc")):
        yield _Cell(tc, row.table)


# -- element details ------------------------------------------------------


def _grid_span(cell) -> int:
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return 1
    span = tc_pr.find(qn("w:gridSpan"))
    if span is None:
        return 1
    try:
        return max(1, int(span.get(qn("w:val"))))
    except (TypeError, ValueError):
        return 1


def _vertical_merge(cell) -> str | None:
    tc_pr = cell._tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    merge = tc_pr.find(qn("w:vMerge"))
    if merge is None:
        return None
    # w:vMerge with no val means "continue the cell above".
    return "restart" if merge.get(qn("w:val")) == "restart" else "continue"


def _is_header_row(row) -> bool:
    tr_pr = row._tr.find(qn("w:trPr"))
    return tr_pr is not None and tr_pr.find(qn("w:tblHeader")) is not None


def _table_style_name(table) -> str | None:
    try:
        return table.style.name if table.style is not None else None
    except (KeyError, AttributeError):
        return None


def _style_name(paragraph) -> str | None:
    try:
        return paragraph.style.name if paragraph.style is not None else None
    except (KeyError, AttributeError):
        return None


def _heading_level(style_name: str | None, paragraph) -> int | None:
    if style_name:
        prefix = "Heading "
        if style_name.startswith(prefix) and style_name[len(prefix):].isdigit():
            return int(style_name[len(prefix):])
        if style_name == "Title":
            return 0

    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        return None
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    try:
        return int(outline.get(qn("w:val"))) + 1
    except (TypeError, ValueError):
        return None


def _list_info(paragraph, context: _ParseContext) -> ListInfo | None:
    numbering = paragraph_numbering(paragraph)
    if numbering is None:
        return None
    num_id, level, from_style = numbering
    marker, num_fmt = context.numbering.next_marker(num_id, level)
    return ListInfo(
        num_id=num_id,
        level=level,
        num_fmt=num_fmt,
        marker=marker,
        from_style=from_style,
    )


def _blocks_to_text(blocks: list[Block]) -> str:
    parts = [
        block.text if isinstance(block, ParagraphBlock) else _table_to_text(block)
        for block in blocks
    ]
    return "\n".join(part for part in parts if part)


def _table_to_text(table: TableBlock) -> str:
    return "\n".join(
        " | ".join(cell.text for cell in row.cells) for row in table.rows
    )
