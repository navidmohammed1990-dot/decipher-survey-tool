"""Character-level formatting extraction."""

from __future__ import annotations

import docx
import pytest

from app.models.document import TextRun
from app.parsing.formatting import (
    extract_runs,
    merge_runs,
    runs_to_text,
    trim_runs_prefix,
)


def _paragraph(*parts):
    """Build a one-paragraph document from (text, bold, italic) triples."""
    document = docx.Document()
    paragraph = document.add_paragraph()
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
    return document, paragraph


def test_extracts_bold_and_italic_per_run():
    _, paragraph = _paragraph(
        ("Plain ", False, False), ("bold", True, False), (" and ", False, False),
        ("italic", False, True),
    )
    runs = extract_runs(paragraph)

    assert [(r.text, r.bold, r.italic) for r in runs] == [
        ("Plain ", False, False),
        ("bold", True, False),
        (" and ", False, False),
        ("italic", False, True),
    ]


def test_bold_and_italic_can_combine_on_one_run():
    _, paragraph = _paragraph(("both", True, True))
    (run,) = extract_runs(paragraph)
    assert run.bold and run.italic


def test_adjacent_runs_with_equal_formatting_are_merged():
    _, paragraph = _paragraph(("Hello ", False, False), ("world", False, False))
    assert [r.text for r in extract_runs(paragraph)] == ["Hello world"]


def test_paragraph_style_bold_is_inherited_by_runs():
    """A run with no direct formatting takes bold from its paragraph style."""
    document = docx.Document()
    style = document.styles["Normal"]
    paragraph = document.add_paragraph("inherited", style=style)
    style.font.bold = True

    (run,) = extract_runs(paragraph)
    assert run.bold is True


def test_direct_formatting_overrides_style():
    document = docx.Document()
    style = document.styles["Normal"]
    style.font.bold = True
    paragraph = document.add_paragraph()
    paragraph.add_run("explicitly not bold").bold = False

    (run,) = extract_runs(paragraph)
    assert run.bold is False


def test_hyperlinked_text_is_not_dropped():
    """Runs inside w:hyperlink are invisible to Paragraph.runs."""
    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("See ")
    paragraph.add_run("link text").bold = True

    # Re-parent the second run under a w:hyperlink element.
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run_el = paragraph._p.findall(qn("w:r"))[1]
    hyperlink = OxmlElement("w:hyperlink")
    paragraph._p.replace(run_el, hyperlink)
    hyperlink.append(run_el)

    assert len(paragraph.runs) == 1, "precondition: python-docx hides the run"
    assert runs_to_text(extract_runs(paragraph)) == "See link text"


def test_tracked_deletions_are_excluded():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    document = docx.Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("kept ")
    paragraph.add_run("removed")

    run_el = paragraph._p.findall(qn("w:r"))[1]
    deletion = OxmlElement("w:del")
    paragraph._p.replace(run_el, deletion)
    deletion.append(run_el)

    assert runs_to_text(extract_runs(paragraph)) == "kept "


def test_merge_runs_drops_empty_runs():
    merged = merge_runs([
        TextRun(text="a"), TextRun(text=""), TextRun(text="b", bold=True),
    ])
    assert [(r.text, r.bold) for r in merged] == [("a", False), ("b", True)]


@pytest.mark.parametrize(
    "n_chars,expected",
    [(0, "Q5. Which"), (4, "Which"), (5, "hich"), (100, "")],
)
def test_trim_runs_prefix_cuts_across_run_boundaries(n_chars, expected):
    runs = [TextRun(text="Q5. "), TextRun(text="Which", bold=True)]
    assert runs_to_text(trim_runs_prefix(runs, n_chars)) == expected


def test_trim_runs_prefix_preserves_formatting_of_remainder():
    runs = [TextRun(text="Q5. Which "), TextRun(text="brands", bold=True)]
    trimmed = trim_runs_prefix(runs, 4)

    assert [(r.text, r.bold) for r in trimmed] == [("Which ", False), ("brands", True)]
