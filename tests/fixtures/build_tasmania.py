"""Questionnaire fixtures for the Phase 6/7 checklists.

Modelled on the Tasmania driver questionnaire that surfaced the Phase 6 bugs:
house style tags each question `ASK ALL, SC`, options live in a two-column
table with explicit codes, and routing text is red.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)


def _red(paragraph, text):
    run = paragraph.add_run(text)
    run.font.color.rgb = RED
    return run


def _options_table(document, rows):
    table = document.add_table(rows=len(rows), cols=2)
    for row, (text, code) in zip(table.rows, rows):
        row.cells[0].text = text
        row.cells[1].text = code
    return table


def build_tasmania(path: Path) -> Path:
    """Q1.1-Q1.6 in the shape that produced the Phase 6 bug report."""
    document = docx.Document()
    document.add_paragraph("Tasmania Driver Survey")

    document.add_paragraph("ASK ALL, NUM")
    document.add_paragraph("Q1.1. What is your age?")
    document.add_paragraph("_________")

    document.add_paragraph("ASK ALL, SC")
    document.add_paragraph("Q1.2. Which of the following best describes your gender identity?")
    _options_table(document, [("Male", "1"), ("Female", "2"), ("Other", "97")])
    _red(document.add_paragraph(), "RANDOMLY ASSIGN OTHER INTO MALE/FEMALE QUOTAS")

    document.add_paragraph("ASK ALL, NUM")
    document.add_paragraph("Q1.3. What is your postcode?")
    document.add_paragraph("_________")

    document.add_paragraph("ASK ALL, SC")
    document.add_paragraph("Q1.4. Do you hold a current Tasmanian driver licence?")
    _options_table(document, [("Yes", "1"), ("No", "2")])

    document.add_paragraph("ASK ALL, OE")
    document.add_paragraph("Q1.5. Why did you answer that way?")
    document.add_paragraph("_________")

    document.add_paragraph("ASK ALL, MC")
    document.add_paragraph("Q1.6. Which of these vehicles do you drive?")
    _options_table(document, [
        ("Car", "1"), ("Motorcycle", "2"), ("Truck", "3"), ("None of these", "99"),
    ])
    _red(document.add_paragraph(), "QUALIFY IF Q1.4 = 1-2 AND Q1.6 <> 99")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def build_format_variations(path: Path) -> Path:
    """The same question written three ways, to prove format independence."""
    document = docx.Document()

    # 1. Two-column table with explicit codes.
    document.add_paragraph("Q1. Which brands have you purchased?")
    _options_table(document, [
        ("Brand A", "1"), ("Brand B", "2"), ("Other, please specify", "97"),
        ("None of these", "99"),
    ])

    # 2. Plain numbered list, codes implied by position.
    document.add_paragraph("Q2. Which brands have you purchased?")
    for text in ("1. Brand A", "2. Brand B", "97. Other, please specify",
                 "99. None of these"):
        document.add_paragraph(text)

    # 3. Bulleted list with parenthetical codes.
    document.add_paragraph("Q3. Which brands have you purchased?")
    for text in ("Brand A (1)", "Brand B (2)", "Other, please specify (97)",
                 "None of these (99)"):
        document.add_paragraph(text, style="List Bullet")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


def build_house_convention(path: Path) -> Path:
    """Routing text with no colour and no keyword the hints recognise.

    The real test of whether the pipeline generalises: nothing about these
    lines is detectable by pattern, so only judgment can place them.
    """
    document = docx.Document()

    document.add_paragraph("Q1. How satisfied are you with your current provider?")
    document.add_paragraph("Please select one.")
    for text in ("Very satisfied", "Satisfied", "Dissatisfied"):
        document.add_paragraph(text)
    # House convention: routing written as a plain, unmarked sentence.
    document.add_paragraph(
        "Respondents choosing Dissatisfied should be shown the follow-up module "
        "before continuing to section 2."
    )

    document.add_paragraph("Q2. How was your current provider chosen?")
    document.add_paragraph("I chose it myself")
    document.add_paragraph("A family member chose it")
    # A genuine option that reads exactly like a randomisation instruction.
    # Planted so the keyword hint and the model must disagree.
    document.add_paragraph("Randomly assigned to me by my employer")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path
