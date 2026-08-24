"""Builds the sample questionnaire used by the test suite.

Generated rather than committed so the fixture stays diffable and its
formatting intent is explicit. Mirrors the worked example in the workflow
document (Q5, brand purchase, select-all) plus a grid and a screener.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIXTURE_PATH = Path(__file__).parent / "sample_questionnaire.docx"


def _runs(paragraph, *parts):
    """Add (text, bold, italic) triples as runs."""
    for text, bold, italic in parts:
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
    return paragraph


def build(path: Path | None = None) -> Path:
    path = Path(path) if path else FIXTURE_PATH
    document = docx.Document()

    document.add_paragraph("Brand Tracker 2026", style="Title")
    intro = document.add_paragraph()
    _runs(intro, ("Thank you for taking part. ", False, False),
                 ("Please answer honestly.", False, True))

    # Screener: single select, options as a bulleted list.
    document.add_paragraph("S1. Which age group do you belong to?")
    _runs(document.add_paragraph(), ("Select one only.", True, False))
    for option in ("Under 18", "18-34", "35-54", "55+"):
        document.add_paragraph(option, style="List Number")

    # Q5: the worked example from the workflow document.
    q5 = document.add_paragraph()
    _runs(q5, ("Q5. Which of the following brands have you ", False, False),
              ("purchased", True, False),
              (" in the last ", False, False),
              ("6 months", False, True),
              ("?", False, False))
    _runs(document.add_paragraph(), ("Please select all that apply.", True, False))
    for option in ("Brand A", "Brand B", "Brand C", "None of these"):
        document.add_paragraph(option, style="List Number")

    # Q6: a grid, expressed as a table between two paragraphs.
    document.add_paragraph("Q6. How much do you agree with each statement?")
    _runs(document.add_paragraph(), ("One response per row.", False, True))

    table = document.add_table(rows=3, cols=3)
    headers = ["Statement", "Agree", "Disagree"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.paragraphs[0].add_run(text).bold = True
    grid_rows = [
        ("The brand is good value", "", ""),
        ("The brand is easy to find", "", ""),
    ]
    for row, values in zip(table.rows[1:], grid_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value

    document.add_paragraph("Ask Q7 only if Q5 = Brand A.", style="Intense Quote")

    # Q7: literal markers typed into the text rather than Word numbering.
    document.add_paragraph("Q7. Why did you choose that brand?")
    for option in ("1. Price", "2. Quality", "3. Availability"):
        document.add_paragraph(option)

    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _runs(closing, ("Thank you.", True, True))

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    return path


if __name__ == "__main__":
    print(build())
