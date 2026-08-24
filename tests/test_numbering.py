"""Word numbering resolution."""

from __future__ import annotations

import docx
import pytest

from app.models.document import ParagraphBlock
from app.parsing.docx_parser import parse_docx
from app.parsing.numbering import _to_letter, _to_roman, format_counter


@pytest.mark.parametrize("value,expected", [(1, "a"), (2, "b"), (26, "z"), (27, "aa"), (28, "ab")])
def test_letter_numbering(value, expected):
    assert _to_letter(value) == expected


@pytest.mark.parametrize("value,expected", [(1, "i"), (4, "iv"), (9, "ix"), (14, "xiv"), (40, "xl")])
def test_roman_numbering(value, expected):
    assert _to_roman(value) == expected


@pytest.mark.parametrize(
    "value,fmt,expected",
    [
        (3, "decimal", "3"),
        (3, "lowerLetter", "c"),
        (3, "upperLetter", "C"),
        (3, "lowerRoman", "iii"),
        (3, "upperRoman", "III"),
        (3, "decimalZero", "03"),
        (3, None, "3"),
        (3, "somethingExotic", "3"),
    ],
)
def test_format_counter(value, fmt, expected):
    assert format_counter(value, fmt) == expected


def test_numbered_options_get_sequential_markers(tmp_path):
    document = docx.Document()
    document.add_paragraph("Q1. Pick one")
    for option in ("Yes", "No", "Maybe"):
        document.add_paragraph(option, style="List Number")
    path = tmp_path / "numbered.docx"
    document.save(path)

    parsed = parse_docx(path)
    markers = [
        b.list_info.marker for b in parsed.blocks
        if isinstance(b, ParagraphBlock) and b.list_info
    ]
    assert markers == ["1.", "2.", "3."]


def test_bulleted_options_report_a_bullet_glyph(tmp_path):
    document = docx.Document()
    document.add_paragraph("Q1. Pick one")
    for option in ("Yes", "No"):
        document.add_paragraph(option, style="List Bullet")
    path = tmp_path / "bulleted.docx"
    document.save(path)

    parsed = parse_docx(path)
    infos = [
        b.list_info for b in parsed.blocks
        if isinstance(b, ParagraphBlock) and b.list_info
    ]
    assert len(infos) == 2
    assert all(info.num_fmt == "bullet" for info in infos)
    assert all(info.marker for info in infos)


def test_paragraphs_without_numbering_have_no_list_info(tmp_path):
    document = docx.Document()
    document.add_paragraph("Q1. Pick one")
    document.add_paragraph("Yes")
    path = tmp_path / "plain.docx"
    document.save(path)

    parsed = parse_docx(path)
    assert all(
        b.list_info is None for b in parsed.blocks if isinstance(b, ParagraphBlock)
    )
