"""Phase 5B — background jobs, progress, cancellation and batching."""

from __future__ import annotations

import threading
import time

import docx
import pytest
from fastapi.testclient import TestClient

from app.classify.jobs import job_manager
from app.classify.ollama import OllamaClient
from app.main import app
from app.parsing.docx_parser import parse_docx
from app.store import draft_store


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as test_client:
        yield test_client


@pytest.fixture(autouse=True)
def _clean():
    draft_store.clear()
    job_manager.clear()
    yield
    job_manager.clear()
    draft_store.clear()


@pytest.fixture(scope="module")
def big_document(tmp_path_factory):
    """A 30-question questionnaire, the size the addendum is about."""
    document = docx.Document()
    for number in range(1, 31):
        document.add_paragraph(f"Q{number}. Question number {number}?")
        document.add_paragraph("Please select one.")
        for option in ("Yes", "No"):
            document.add_paragraph(option)
    path = tmp_path_factory.mktemp("big") / "big.docx"
    document.save(path)
    return parse_docx(path).model_dump(mode="json")


def install_ai(monkeypatch, delay=0.0):
    """A model that answers instantly, or slowly enough to observe progress."""
    from app.api import routes_classify

    class Fake(OllamaClient):
        def generate_json(self, system, prompt):
            if delay:
                time.sleep(delay)
            return {"element": "radio", "title_lines": [0], "comment_lines": [1],
                    "option_lines": [2, 3], "confidence": 0.9, "notes": "ok"}

        def is_available(self):
            return True

    monkeypatch.setattr(routes_classify, "build_client", Fake)


def wait_for(client, job_id, states=("completed", "cancelled", "failed"), timeout=20):
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        status = client.get(f"/api/classify/{job_id}/status").json()
        if status["state"] in states:
            return status
        time.sleep(0.05)
    raise AssertionError(f"job {job_id} did not finish: {status}")


def start(client, document, **body):
    return client.post("/api/classify/start", json={"document": document, **body})


# -- starting a job -------------------------------------------------------


def test_start_returns_immediately_with_a_job_id(client, big_document, monkeypatch):
    install_ai(monkeypatch, delay=0.05)
    body = start(client, big_document).json()

    assert len(body["job_id"]) > 0
    assert body["total"] == 30
    assert body["labels"][:2] == ["Q1", "Q2"]
    assert body["slow_warning_seconds"] > 0


def test_status_reports_real_progress(client, big_document, monkeypatch):
    """Checklist: progress reflects real elapsed/estimated time."""
    install_ai(monkeypatch, delay=0.03)
    job_id = start(client, big_document).json()["job_id"]

    seen = []
    deadline = time.monotonic() + 20
    while time.monotonic() < deadline:
        status = client.get(f"/api/classify/{job_id}/status").json()
        seen.append(status["completed"])
        if status["state"] != "running":
            break
        time.sleep(0.05)

    assert seen == sorted(seen), "progress must never go backwards"
    assert max(seen) == 30

    final = client.get(f"/api/classify/{job_id}/status").json()
    assert final["state"] == "completed"
    assert final["elapsed_seconds"] > 0


def test_estimate_is_absent_until_something_finishes(client, big_document, monkeypatch):
    """No fabricated ETA: there is nothing to extrapolate from at zero."""
    install_ai(monkeypatch, delay=0.2)
    job_id = start(client, big_document).json()["job_id"]

    first = client.get(f"/api/classify/{job_id}/status").json()
    if first["completed"] == 0:
        assert first["estimated_remaining_seconds"] is None

    wait_for(client, job_id)


def test_estimate_projects_from_the_running_average(client, big_document, monkeypatch):
    install_ai(monkeypatch, delay=0.05)
    job_id = start(client, big_document).json()["job_id"]

    estimate = None
    deadline = time.monotonic() + 20
    while time.monotonic() < deadline:
        status = client.get(f"/api/classify/{job_id}/status").json()
        if status["completed"] >= 3 and status["state"] == "running":
            estimate = status["estimated_remaining_seconds"]
            break
        if status["state"] != "running":
            break
        time.sleep(0.02)

    if estimate is not None:
        # 30 questions at ~50ms each: seconds, not minutes or milliseconds.
        assert 0 < estimate < 60

    wait_for(client, job_id)


def test_a_finished_job_reports_no_remaining_time(client, big_document, monkeypatch):
    install_ai(monkeypatch)
    job_id = start(client, big_document).json()["job_id"]
    final = wait_for(client, job_id)

    assert final["estimated_remaining_seconds"] is None
    assert final["completed"] == final["total"]


def test_the_slow_flag_trips_on_a_long_projection(client, big_document, monkeypatch):
    """Checklist: warn when the projection exceeds the threshold."""
    from app.api import routes_classify
    from dataclasses import replace

    monkeypatch.setattr(
        routes_classify, "settings", replace(routes_classify.settings, slow_classify_seconds=0.01)
    )
    install_ai(monkeypatch, delay=0.05)
    job_id = start(client, big_document).json()["job_id"]

    tripped = False
    deadline = time.monotonic() + 20
    while time.monotonic() < deadline:
        status = client.get(f"/api/classify/{job_id}/status").json()
        tripped = tripped or status["slow"]
        if status["state"] != "running":
            break
        time.sleep(0.02)

    assert tripped, "a 30-question run must exceed a 10ms threshold"


def test_unknown_job_is_a_404(client):
    assert client.get("/api/classify/nope/status").status_code == 404
    assert client.post("/api/classify/nope/cancel").status_code == 404


# -- cancellation ---------------------------------------------------------


def test_cancelling_keeps_completed_questions(client, big_document, monkeypatch):
    """Checklist: partial results stay usable, not thrown away."""
    install_ai(monkeypatch, delay=0.05)
    job_id = start(client, big_document).json()["job_id"]

    # Let a few land, then stop.
    deadline = time.monotonic() + 20
    while time.monotonic() < deadline:
        if client.get(f"/api/classify/{job_id}/status").json()["completed"] >= 3:
            break
        time.sleep(0.02)

    client.post(f"/api/classify/{job_id}/cancel")
    final = wait_for(client, job_id, states=("cancelled", "completed"))

    kept = client.get("/api/questions").json()["questions"]
    assert len(kept) >= 3
    assert len(kept) == final["completed"]
    assert all(q["element"] for q in kept)


def test_a_cancelled_job_stops_early(client, big_document, monkeypatch):
    install_ai(monkeypatch, delay=0.05)
    job_id = start(client, big_document).json()["job_id"]

    deadline = time.monotonic() + 20
    while time.monotonic() < deadline:
        if client.get(f"/api/classify/{job_id}/status").json()["completed"] >= 2:
            break
        time.sleep(0.02)

    client.post(f"/api/classify/{job_id}/cancel")
    final = wait_for(client, job_id, states=("cancelled",))

    assert final["state"] == "cancelled"
    assert final["completed"] < 30


def test_cancelled_work_can_still_be_generated(client, big_document, monkeypatch):
    install_ai(monkeypatch, delay=0.04)
    job_id = start(client, big_document).json()["job_id"]

    deadline = time.monotonic() + 20
    while time.monotonic() < deadline:
        if client.get(f"/api/classify/{job_id}/status").json()["completed"] >= 2:
            break
        time.sleep(0.02)
    client.post(f"/api/classify/{job_id}/cancel")
    wait_for(client, job_id, states=("cancelled", "completed"))

    generated = client.post("/api/generate", json={}).json()
    assert generated["well_formed"] is True
    assert generated["question_count"] >= 2


# -- batching -------------------------------------------------------------


def test_a_batch_classifies_only_the_selected_labels(client, big_document, monkeypatch):
    install_ai(monkeypatch)
    job_id = start(client, big_document, labels=["Q1", "Q2", "Q3"]).json()["job_id"]
    wait_for(client, job_id)

    labels = [q["label"] for q in client.get("/api/questions").json()["questions"]]
    assert labels == ["Q1", "Q2", "Q3"]


def test_two_batches_merge_into_one_review_set(client, big_document, monkeypatch):
    """Checklist: Q1-Q15 then Q16-Q30 merges into 30, not overwrites."""
    install_ai(monkeypatch)

    first = [f"Q{n}" for n in range(1, 16)]
    wait_for(client, start(client, big_document, labels=first).json()["job_id"])
    assert len(client.get("/api/questions").json()["questions"]) == 15

    second = [f"Q{n}" for n in range(16, 31)]
    wait_for(client, start(client, big_document, labels=second).json()["job_id"])

    body = client.get("/api/questions").json()
    assert len(body["questions"]) == 30
    assert body["summary"]["total"] == 30


def test_merged_batches_keep_document_order(client, big_document, monkeypatch):
    """A late batch sorts into place rather than appending."""
    install_ai(monkeypatch)
    wait_for(client, start(client, big_document, labels=["Q20", "Q21"]).json()["job_id"])
    wait_for(client, start(client, big_document, labels=["Q1", "Q2"]).json()["job_id"])

    labels = [q["label"] for q in client.get("/api/questions").json()["questions"]]
    assert labels == ["Q1", "Q2", "Q20", "Q21"]


def test_reclassifying_a_label_replaces_it_without_duplicating(client, big_document, monkeypatch):
    install_ai(monkeypatch)
    wait_for(client, start(client, big_document, labels=["Q1"]).json()["job_id"])
    wait_for(client, start(client, big_document, labels=["Q1"]).json()["job_id"])

    labels = [q["label"] for q in client.get("/api/questions").json()["questions"]]
    assert labels == ["Q1"]


def test_programmer_edits_survive_a_later_batch(client, big_document, monkeypatch):
    """Classifying Q2 must not undo a correction already made to Q1."""
    install_ai(monkeypatch)
    wait_for(client, start(client, big_document, labels=["Q1"]).json()["job_id"])
    client.patch("/api/questions/Q1", json={"element": "textarea", "dev_notes": "checked"})

    wait_for(client, start(client, big_document, labels=["Q2"]).json()["job_id"])

    questions = {q["label"]: q for q in client.get("/api/questions").json()["questions"]}
    assert questions["Q1"]["element"] == "textarea"
    assert questions["Q1"]["dev_notes"] == "checked"


def test_a_different_questionnaire_starts_a_fresh_set(client, big_document, sample_docx, monkeypatch):
    """Merging is per-document; a new upload must not inherit the old one."""
    install_ai(monkeypatch)
    wait_for(client, start(client, big_document, labels=["Q1", "Q2"]).json()["job_id"])

    other = parse_docx(sample_docx).model_dump(mode="json")
    wait_for(client, start(client, other, labels=["Q5"]).json()["job_id"])

    labels = [q["label"] for q in client.get("/api/questions").json()["questions"]]
    assert labels == ["Q5"]


def test_unknown_labels_are_rejected(client, big_document, monkeypatch):
    install_ai(monkeypatch)
    response = start(client, big_document, labels=["Q1", "Q999"])

    assert response.status_code == 400
    assert "Q999" in response.json()["detail"]


def test_a_document_with_no_questions_is_rejected(client, monkeypatch, tmp_path):
    install_ai(monkeypatch)
    path = tmp_path / "empty.docx"
    docx.Document().save(path)

    response = start(client, parse_docx(path).model_dump(mode="json"))
    assert response.status_code == 400
    assert "no questions" in response.json()["detail"]


# -- degradation ----------------------------------------------------------


def test_an_offline_model_still_completes_the_job(client, big_document, monkeypatch):
    from app.api import routes_classify

    monkeypatch.setattr(
        routes_classify, "build_client",
        lambda: OllamaClient(base_url="http://127.0.0.1:1", timeout=1.0),
    )
    job_id = start(client, big_document, labels=["Q1", "Q2"]).json()["job_id"]
    final = wait_for(client, job_id, timeout=30)

    assert final["state"] == "completed"
    assert final["fallback_count"] == 2
    assert all(q["needs_review"] for q in client.get("/api/questions").json()["questions"])


def test_concurrent_jobs_do_not_interfere(client, big_document, monkeypatch):
    install_ai(monkeypatch, delay=0.02)
    first = start(client, big_document, labels=[f"Q{n}" for n in range(1, 6)]).json()["job_id"]
    second = start(client, big_document, labels=[f"Q{n}" for n in range(6, 11)]).json()["job_id"]

    wait_for(client, first)
    wait_for(client, second)

    labels = [q["label"] for q in client.get("/api/questions").json()["questions"]]
    assert labels == [f"Q{n}" for n in range(1, 11)]
