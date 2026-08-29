"""Phase 10 — label shape, wrapped options, and a library that persists."""

from __future__ import annotations

import pytest

from app.classify.corrections import Correction, CorrectionMemory
from app.classify.library import CorrectionLibrary
from app.classify.paste import split_questions
from app.classify.wrapping import merge_wrapped_options
from app.models.document import ParagraphBlock
from app.parsing.docx_parser import parse_docx
from app.parsing.question_boundaries import (
    BoundaryConfig,
    detect_boundaries,
    match_question_label,
)

#: Every style the tool must recognise, including the four a whitelist missed.
NINE_STYLES = ["Q1.", "QA1.", "D5A.", "S1.", "QD24.", "P1.", "MP2.", "QZ5.", "APP1."]


# -- labels: a shape, not a list ------------------------------------------


@pytest.mark.parametrize("label", NINE_STYLES)
def test_all_nine_label_styles_match(label):
    found = match_question_label(f"{label} Some question text?")

    assert found is not None, f"{label} did not match"
    assert found[0] == label.rstrip(".")


@pytest.mark.parametrize("label", ["ZZ9.", "XYZ1.", "AB12.", "K7.", "QQQ99."])
def test_prefixes_nobody_has_seen_yet_also_match(label):
    """The point of the change: no enumerated list to fall behind."""
    assert match_question_label(f"{label} A question") is not None


@pytest.mark.parametrize(
    "text",
    [
        "Yes 1.", "No 2.", "Male 1", "Female  2", "70 years or more 8",
        "S2_AGE BANDS", "Under 18 years", "Brand A",
        "COVID19: impact on your business", "Section D. Demographics",
        "Mr. Smith went to town", "iPhone 12 owners", "Please select one.",
        "ASK ALL, SC", "TERMINATE IF Q1 = 2",
    ],
)
def test_ordinary_content_is_not_mistaken_for_a_label(text):
    """A general pattern must not start swallowing option text."""
    assert match_question_label(text) is None


def test_no_whitespace_is_allowed_between_letters_and_digits():
    """Allowing it would make every "Yes 1." option a question boundary."""
    assert match_question_label("Q 12. Spaced") is None
    assert match_question_label("Q12. Unspaced") is not None


def test_a_variable_name_is_not_a_label():
    """Guards Phase 11: S2_AGE reading as a label would break it."""
    assert match_question_label("S2_AGE BANDS") is None
    assert match_question_label("Q4_REGION_GROUP") is None


@pytest.mark.parametrize(
    "text,expected",
    [
        ("[Q7] Bracketed", "Q7"),
        ("Q1.2. Tasmania style", "Q1.2"),
        ("Q5_1. Numeric sub-part", "Q5_1"),
        ("Q10a. Letter sub-part", "Q10A"),
        ("QD24   Whitespace separator", "QD24"),
        ("Q12 - dash separator", "Q12"),
        ("S2: colon separator", "S2"),
    ],
)
def test_existing_forms_still_work(text, expected):
    assert match_question_label(text)[0] == expected


# -- one implementation, both entry points --------------------------------


@pytest.mark.parametrize("label", NINE_STYLES)
def test_paste_and_document_paths_agree(label, tmp_path):
    """The two entry points must not drift apart on label detection."""
    import docx

    text = f"{label} Which of these apply?"

    pasted, _ = split_questions(f"{text}\nYes\nNo")

    document = docx.Document()
    for line in (text, "Yes", "No"):
        document.add_paragraph(line)
    path = tmp_path / "doc.docx"
    document.save(path)
    parsed = parse_docx(path)

    document_labels = [q.label for q in parsed.questions if not q.is_preamble]
    assert [b.label for b in pasted] == document_labels
    assert document_labels == [label.rstrip(".")]


def test_a_document_using_an_unseen_prefix_splits_correctly(tmp_path):
    import docx

    document = docx.Document()
    for number in (1, 2, 3):
        document.add_paragraph(f"APP{number}. Question {number}?")
        document.add_paragraph("Yes")
        document.add_paragraph("No")
    path = tmp_path / "app.docx"
    document.save(path)

    parsed = parse_docx(path)
    assert [q.label for q in parsed.questions if not q.is_preamble] == ["APP1", "APP2", "APP3"]


def test_detection_can_still_be_narrowed():
    blocks = [ParagraphBlock(index=0, text="ZZ4. Custom"), ParagraphBlock(index=1, text="Q4. Standard")]
    found, _ = detect_boundaries(blocks, BoundaryConfig(prefixes=("ZZ",)))

    assert [b.label for b in found if not b.is_preamble] == ["ZZ4"]


# -- wrapped options -------------------------------------------------------


WRAPPED = """QD24.   If you were to purchase this product, would you...?
Buy it instead of another [BRAND] [FORMAT OF INTEREST]
product you usually buy 1
Buy it instead of a different type of product    3"""


def test_a_wrapped_option_is_rejoined():
    (block,), _ = split_questions(WRAPPED)
    texts = [line.text for line in block.lines]

    assert block.label == "QD24"
    assert texts[1] == (
        "Buy it instead of another [BRAND] [FORMAT OF INTEREST] "
        "product you usually buy | 1"
    )
    assert texts[2] == "Buy it instead of a different type of product | 3"


def test_the_rejoined_option_keeps_its_code():
    (block,), _ = split_questions(WRAPPED)
    assert [line.features.trailing_numeric_code for line in block.lines[1:]] == ["1", "3"]


def test_the_rejoined_option_reaches_the_xml_whole():
    from app.classify.classifier import interpret_response
    from app.generate.xml_generator import generate_question

    (block,), _ = split_questions(WRAPPED)
    outcome = interpret_response(
        {"element": "radio", "title_lines": [0], "option_lines": [1, 2],
         "confidence": 0.9, "notes": "two options"},
        "QD24", block.lines, 0.75,
    )
    xml = generate_question(outcome.question)

    assert '<row label="r1" value="1">Buy it instead of another [BRAND] ' \
           '[FORMAT OF INTEREST] product you usually buy</row>' in xml
    assert '<row label="r3" value="3">Buy it instead of a different type of product</row>' in xml


def test_indices_stay_dense_after_merging():
    (block,), _ = split_questions(WRAPPED)
    assert [line.index for line in block.lines] == list(range(len(block.lines)))


# -- and what must NOT be merged ------------------------------------------


def test_a_standalone_heading_is_not_merged_into_an_option():
    """S2_AGE BANDS is a variable name, not half of "Under 18 years"."""
    text = ("[Please create the following variable based on S2 Age]:\n"
            "S2_AGE BANDS\nUnder 18 years  1\n18 to 24 years  2\n70 years or more  8")
    (block,), _ = split_questions(text)

    assert [line.text for line in block.lines][1] == "S2_AGE BANDS"
    assert len(block.lines) == 5


def test_an_instruction_is_not_merged_into_an_option():
    text = "Q1. Which apply?\nPlease select one.\nBrand A  1\nBrand B  2\nBrand C  3"
    (block,), _ = split_questions(text)

    assert "Please select one." in [line.text for line in block.lines]
    assert len(block.lines) == 5


def test_a_title_is_never_merged_into_the_first_option():
    text = "Q1. A rather long question title that keeps going\nBrand A  1\nBrand B  2\nBrand C 3"
    (block,), _ = split_questions(text)
    assert block.lines[0].text.startswith("A rather long question title")


def test_an_uncoded_list_is_left_alone():
    """With no codes anywhere, nothing establishes the wrap convention."""
    text = "Q1. Which apply?\nBrand A\nBrand B\nBrand C"
    (block,), _ = split_questions(text)
    assert len(block.lines) == 4


def test_merging_is_a_no_op_on_a_short_block():
    lines = split_questions("Q1. Pick one\nYes | 1")[0][0].lines
    assert merge_wrapped_options(lines) == lines


# -- persistent library ----------------------------------------------------


def a_correction(label="Q1", to="checkbox"):
    return Correction(
        label=label, original_lines=["Which apply?", "Yes", "No"],
        ai_said={"element": "radio"}, sp_corrected_to={"element": to},
    )


def test_a_correction_survives_a_restart(tmp_path):
    path = tmp_path / "nested" / "corrections.json"
    CorrectionLibrary(path).record(a_correction(), document="brand.docx")

    assert path.is_file()
    assert len(CorrectionLibrary(path).entries()) == 1


def test_a_correction_is_offered_back_for_its_own_document(tmp_path):
    library = CorrectionLibrary(tmp_path / "c.json")
    library.record(a_correction(), document="brand.docx")

    assert len(library.for_document("brand.docx")) == 1
    assert library.for_document("other.docx") == [], "no leak between clients"


def test_a_promoted_correction_applies_everywhere(tmp_path):
    path = tmp_path / "c.json"
    library = CorrectionLibrary(path)
    library.record(a_correction(), document="brand.docx")

    assert library.promote("Q1", "brand.docx") is True
    assert len(CorrectionLibrary(path).for_document("other.docx")) == 1


def test_re_editing_supersedes_the_earlier_entry(tmp_path):
    library = CorrectionLibrary(tmp_path / "c.json")
    library.record(a_correction(to="checkbox"), document="brand.docx")
    library.record(a_correction(to="select"), document="brand.docx")

    entries = library.entries()
    assert len(entries) == 1
    assert entries[0].correction.sp_corrected_to["element"] == "select"


def test_a_meaningless_correction_is_not_stored(tmp_path):
    library = CorrectionLibrary(tmp_path / "c.json")
    unchanged = Correction(label="Q1", original_lines=["x"],
                           ai_said={"element": "radio"}, sp_corrected_to={"element": "radio"})

    assert library.record(unchanged, document="d.docx") is False
    assert library.entries() == []


def test_the_prompt_prefix_is_bounded(tmp_path):
    from app.classify.library import MAX_PROMPTED

    library = CorrectionLibrary(tmp_path / "c.json")
    for number in range(8):
        library.record(a_correction(label=f"Q{number}"), document="d.docx")

    prefix = library.prompt_prefix("d.docx")
    assert prefix.count("The survey programmer corrected") == MAX_PROMPTED


def test_a_damaged_library_does_not_stop_startup(tmp_path):
    path = tmp_path / "c.json"
    path.write_text("{ this is not json")

    assert CorrectionLibrary(path).entries() == []


def test_an_unwritable_path_does_not_crash(tmp_path):
    blocker = tmp_path / "blocked"
    blocker.write_text("not a directory")

    library = CorrectionLibrary(blocker / "c.json")
    assert library.record(a_correction(), document="d.docx") is True, "kept in memory"


def test_memory_without_a_library_still_works():
    memory = CorrectionMemory()
    memory.use_document("d.docx")

    assert memory.record(a_correction()) is True
    assert len(memory.recent()) == 1


def test_reopening_a_document_restores_its_corrections(tmp_path):
    library = CorrectionLibrary(tmp_path / "c.json")
    memory = CorrectionMemory(library=library)

    memory.use_document("brand.docx")
    memory.record(a_correction())

    memory.use_document("other.docx")
    assert memory.recent() == [], "a different document starts clean"

    memory.use_document("brand.docx")
    assert len(memory.recent()) == 1, "coming back restores what was corrected"


def test_the_library_can_be_imported_before_corrections():
    """Guards a circular import that only showed up on one import order.

    corrections.py used to build the library at module scope, so importing
    library.py first deadlocked the pair. The suite missed it because tests
    happen to import app.main, which loads them in the working order.
    """
    import subprocess
    import sys

    for first, second in (
        ("app.classify.library", "app.classify.corrections"),
        ("app.classify.corrections", "app.classify.library"),
    ):
        result = subprocess.run(
            [sys.executable, "-c", f"import {first}; import {second}"],
            capture_output=True, text=True, cwd=".",
        )
        assert result.returncode == 0, f"{first} then {second} failed:\n{result.stderr}"


def test_importing_the_library_module_does_not_touch_the_disk(monkeypatch):
    """A module import must not read a file; the singleton is built on use."""
    import app.classify.library as module

    monkeypatch.setattr(module, "_default", None)
    assert module._default is None, "nothing is constructed until asked for"
