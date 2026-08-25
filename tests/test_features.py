"""Phase 7 feature extraction — observations, never decisions."""

from __future__ import annotations

import pytest

from app.classify.features import (
    LineFeatures,
    color_name,
    detect_trailing_code,
    detect_type_tag,
    extract_features,
    matches_routing_keyword,
    strip_trailing_code,
)
from app.models.document import TextRun


def runs(text, **kwargs):
    return [TextRun(text=text, **kwargs)]


# -- the architectural guarantee ------------------------------------------


def test_no_field_assigns_a_role():
    """A hint may describe a line; it may never label one.

    This is the whole correction in the revised brief: patterns observe,
    judgment decides. A field named `is_routing` here would be the old design
    creeping back.
    """
    forbidden = {"role", "is_routing", "is_option", "is_title", "is_comment",
                 "routing_instruction", "line_role"}
    assert forbidden.isdisjoint(LineFeatures.model_fields)


def test_hinted_lines_are_still_offered_to_the_classifier(tmp_path):
    """Nothing is filtered out by its hints — the classifier sees every line.

    Under the old design a routing keyword removed a line before the model
    saw it. Here the strongest possible hints must change nothing but the
    evidence attached.
    """
    import docx

    from app.classify.lines import question_lines
    from app.parsing.docx_parser import parse_docx

    document = docx.Document()
    document.add_paragraph("Q1. Which brands?")
    document.add_paragraph("TERMINATE IF Q1 = 2")
    document.add_paragraph("ASK ALL, SC")
    document.add_paragraph("Brand A")
    path = tmp_path / "hinted.docx"
    document.save(path)

    parsed = parse_docx(path)
    question = next(q for q in parsed.questions if not q.is_preamble)
    lines = question_lines(parsed, question)

    assert [line.text for line in lines] == [
        "Which brands?", "TERMINATE IF Q1 = 2", "ASK ALL, SC", "Brand A",
    ]
    assert [line.index for line in lines] == [0, 1, 2, 3]
    assert lines[1].features.matches_routing_keyword is True


# -- individual observations ----------------------------------------------


@pytest.mark.parametrize(
    "text,expected",
    [
        ("Male | 1", "1"),
        ("Other | 97", "97"),
        ("Other (97)", "97"),
        ("None of these [99]", "99"),
        ("Brand A (1)", "1"),
    ],
)
def test_trailing_codes_are_detected(text, expected):
    assert detect_trailing_code(text) == expected


@pytest.mark.parametrize("text", ["Under 18", "55+", "Aged 18 to 34", "Brand 500", "2 or more"])
def test_a_bare_trailing_number_is_not_a_code(text):
    """"Under 18" must not become code 18 — a separator or bracket is required."""
    assert detect_trailing_code(text) is None


@pytest.mark.parametrize(
    "text,expected",
    [("Male | 1", "Male"), ("Other (97)", "Other"), ("Under 18", "Under 18")],
)
def test_codes_are_stripped_from_option_text(text, expected):
    assert strip_trailing_code(text) == expected


@pytest.mark.parametrize(
    "text,expected",
    [
        ("ASK ALL, SC", "SC"),
        ("ASK ALL, MC", "MC"),
        ("ASK ALL, OE", "OE"),
        ("[SC]", "SC"),
        ("ASK IF Q1 = 1, NUM", "NUM"),
    ],
)
def test_type_tags_are_detected(text, expected):
    assert detect_type_tag(text) == expected


@pytest.mark.parametrize("text", ["Scotland", "MCDONALDS", "Which one?", "Occasionally"])
def test_words_containing_tag_letters_are_not_type_tags(text):
    assert detect_type_tag(text) is None


@pytest.mark.parametrize(
    "text",
    ["ASK ALL", "TERMINATE IF Q1 = 2", "SKIP TO Q9", "QUALIFY IF Q1.4 = 1-2",
     "RANDOMLY ASSIGN OTHER INTO QUOTAS", "PROGRAMMER NOTE: rotate", "randomize options"],
)
def test_routing_keywords_are_detected(text):
    assert matches_routing_keyword(text) is True


@pytest.mark.parametrize(
    "text",
    ["Which brands have you purchased?", "Please select all that apply.",
     "None of these", "I would ask a friend"],
)
def test_ordinary_text_matches_no_routing_keyword(text):
    assert matches_routing_keyword(text) is False


@pytest.mark.parametrize(
    "hex_value,expected",
    [("FF0000", "red"), ("C00000", "red"), ("0070C0", "blue"),
     ("000000", None), ("1A1A1A", None), (None, None)],
)
def test_colour_names(hex_value, expected):
    assert color_name(hex_value) == expected


def test_bold_is_observed_only_when_the_whole_line_is_bold():
    mixed = [TextRun(text="Some ", bold=True), TextRun(text="words")]
    assert extract_features("Some words", mixed).is_bold is False
    assert extract_features("Some words", runs("Some words", bold=True)).is_bold is True


def test_a_red_line_is_only_reported_as_red():
    """Colour is evidence. It does not make the line an instruction."""
    line = extract_features("Brand A", runs("Brand A", color="FF0000"))

    assert line.is_colored is True
    assert line.color_hint == "red"
    assert line.matches_routing_keyword is False


def test_hints_render_only_what_was_observed():
    assert extract_features("Which brands?", runs("Which brands?")).as_prompt_hints() == "none"

    rendered = extract_features(
        "RANDOMLY ASSIGN OTHER", runs("RANDOMLY ASSIGN OTHER", color="FF0000")
    ).as_prompt_hints()
    assert "color_hint=red" in rendered
    assert "matches_routing_keyword=true" in rendered
